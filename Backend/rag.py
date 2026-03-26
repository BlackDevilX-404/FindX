import os
import json
import re
import time
import uuid
import csv
from dataclasses import dataclass, replace
from pathlib import Path
from typing import Any, Optional

import chromadb
import fitz
from dotenv import load_dotenv
from huggingface_hub.utils import disable_progress_bars
from huggingface_hub.utils import logging as hf_logging
from langchain_core.messages import HumanMessage, SystemMessage
from langchain_groq import ChatGroq
from openpyxl import load_workbook
from pydantic import BaseModel, Field
from sentence_transformers import SentenceTransformer
from transformers.utils import logging as transformers_logging
from docx import Document as DocxDocument

os.environ.setdefault("HF_HUB_DISABLE_PROGRESS_BARS", "1")
hf_logging.set_verbosity_error()
transformers_logging.set_verbosity_error()
disable_progress_bars()
transformers_logging.disable_progress_bar()

try:
    from .auth import ROLE_ADMIN, ROLE_DEVELOPER, ROLE_HR
    from .db import list_document_records
    from .pdf_chroma_ingest import ChromaMultimodalDB
    from .process_ppt import Ppt2Pdf
except ImportError:
    from auth import ROLE_ADMIN, ROLE_DEVELOPER, ROLE_HR
    from db import list_document_records
    from pdf_chroma_ingest import ChromaMultimodalDB
    from process_ppt import Ppt2Pdf

load_dotenv()
load_dotenv(os.path.join(os.path.dirname(__file__), ".env"))

ROLE_CATEGORY_ACCESS = {
    ROLE_ADMIN: ["HR", "TECH", "FINANCE", "GENERAL"],
    ROLE_HR: ["HR", "GENERAL"],
    ROLE_DEVELOPER: ["TECH", "GENERAL"],
}
ROLE_VISIBILITY_ACCESS = {
    ROLE_ADMIN: ["private", "hr", "developer", "both"],
    ROLE_HR: ["hr", "both"],
    ROLE_DEVELOPER: ["developer", "both"],
}
VALID_CATEGORIES = sorted({category for values in ROLE_CATEGORY_ACCESS.values() for category in values})
CATEGORY_VARIANTS = {
    category: {category, category.lower(), category.title()}
    for category in VALID_CATEGORIES
}
VALID_VISIBILITY_SCOPES = ["private", "hr", "developer", "both"]

WORD_PATTERN = re.compile(r"[a-zA-Z0-9]+")
STOP_WORDS = {
    "a",
    "an",
    "and",
    "are",
    "about",
    "as",
    "at",
    "be",
    "by",
    "document",
    "file",
    "for",
    "from",
    "how",
    "in",
    "is",
    "it",
    "of",
    "on",
    "or",
    "pdf",
    "rule",
    "rules",
    "say",
    "says",
    "tell",
    "the",
    "to",
    "what",
    "when",
    "where",
    "who",
    "why",
    "with",
}

MIN_HIGH_CONFIDENT_SCORE = 0.60
MIN_CONFIDENT_SCORE = 0.30
MIN_LEXICAL_OVERLAP = 0.08
MIN_LIGHT_OVERLAP = 0.04
FOLLOW_UP_HINTS = {
    "it",
    "its",
    "they",
    "them",
    "that",
    "those",
    "this",
    "these",
    "he",
    "she",
    "their",
    "there",
    "same",
}
SMALL_TALK_PATTERNS = [
    re.compile(r"^(hi|hello|hey)( there)?[!. ]*$", re.IGNORECASE),
    re.compile(r"^good (morning|afternoon|evening)[!. ]*$", re.IGNORECASE),
    re.compile(r"^(thanks|thank you|ok thanks)[!. ]*$", re.IGNORECASE),
    re.compile(r"^(who are you|what can you do)\??$", re.IGNORECASE),
]
FILE_REFERENCE_PATTERN = re.compile(r"\b[\w.-]+\.(pdf|ppt|pptx|docx|txt|md|csv|json|xlsx)\b", re.IGNORECASE)
FILE_TYPE_HINT_PATTERN = re.compile(r"\b(pdf|ppt|pptx|docx|txt|md|csv|json|xlsx)\b", re.IGNORECASE)
PAGE_REFERENCE_PATTERN = re.compile(r"\bpage\s*(?:#|no\.?|number)?\s*(\d{1,4})\b", re.IGNORECASE)
DEFAULT_AGENT_TOP_K = 4
MIN_AGENT_TOP_K = 2
MAX_AGENT_TOP_K = 8
MAX_AGENT_STEPS = 3
MAX_VERIFICATION_REFINEMENTS = 1
TARGET_CHUNK_TOKENS = 280
CHUNK_OVERLAP_TOKENS = 60
SHORT_PARAGRAPH_WORDS = 45
MAX_EXPANDED_CONTEXT_TOKENS = 360
EMBEDDING_BATCH_SIZE = 32
INGEST_WRITE_BATCH_SIZE = 96
INGEST_PROGRESS_PAGE_INTERVAL = 10


class SourceItem(BaseModel):
    id: str
    document: str
    doc_uuid: str
    snippet: str
    page: int | None = None
    confidence: str | None = None


class QueryResult(BaseModel):
    answer: str
    explanation: str
    sources: list[SourceItem] = Field(default_factory=list)


class IngestResult(BaseModel):
    document_id: str
    document: str
    category: str
    sensitivity: str | None = None
    chunks_indexed: int


@dataclass
class RetrievedChunk:
    id: str
    document_id: str
    document: str
    category: str
    sensitivity: str | None
    page: Optional[int]
    paragraph: Optional[int]
    window: Optional[int]
    text: str
    score: float
    semantic_score: float
    lexical_score: float
    retrieval_query: str


@dataclass
class PlannedQuery:
    text: str
    reason: str


@dataclass
class RetrievalAction:
    action: str
    query: str
    top_k: int
    document_id: Optional[str]
    reason: str


@dataclass
class GapQuery:
    text: str
    reason: str
    keep_document_scope: bool = True


@dataclass
class VerificationResult:
    grounded: bool
    needs_more_retrieval: bool
    reason: str
    gap_query: Optional[GapQuery] = None


@dataclass
class AnswerAttempt:
    answer: str
    chunks: list[RetrievedChunk]
    explanation: str
    sources: list[SourceItem]
    verification: Optional[VerificationResult] = None
    refinement_count: int = 0


@dataclass
class ResolvedQueryContext:
    query_text: str
    role: str
    chat_id: str | None
    normalized_history: list[dict[str, str]]
    resolved_doc_uuid: str | None
    requested_page: int | None
    top_k: int


def _normalize_unicode(text: str) -> str:
    return (
        (text or "")
        .replace("\u2022", "- ")
        .replace("\u25cf", "- ")
        .replace("\u2013", "-")
        .replace("\u2014", "-")
        .replace("\u00a0", " ")
        .replace("\u200b", "")
    )


def normalize_text(text: str) -> str:
    normalized = _normalize_unicode(text)
    return re.sub(r"\s+", " ", normalized).strip()


def normalize_structured_text(text: str) -> str:
    normalized = _normalize_unicode(text).replace("\r\n", "\n").replace("\r", "\n")
    blocks: list[str] = []

    for raw_block in re.split(r"\n{2,}", normalized):
        lines = [re.sub(r"[ \t]+", " ", line).strip() for line in raw_block.split("\n")]
        block_text = " ".join(line for line in lines if line)
        if block_text:
            blocks.append(block_text)

    return "\n\n".join(blocks).strip()


def split_paragraphs(text: str) -> list[str]:
    structured = normalize_structured_text(text)
    paragraphs = [normalize_text(paragraph) for paragraph in re.split(r"\n{2,}", structured) if paragraph.strip()]
    if paragraphs:
        return paragraphs

    normalized = normalize_text(text)
    return [normalized] if normalized else []


def sentence_chunks(text: str) -> list[str]:
    return re.split(r"(?<=[.!?])\s+", text)


def sliding_chunks(
    tokens: list[str],
    size: int = TARGET_CHUNK_TOKENS,
    overlap: int = CHUNK_OVERLAP_TOKENS,
):
    index = 0
    while index < len(tokens):
        yield tokens[index : index + size]
        index += size - overlap


def extract_keywords(text: str) -> set[str]:
    return {
        token
        for token in WORD_PATTERN.findall((text or "").lower())
        if token not in STOP_WORDS and len(token) > 2
    }


def keyword_overlap(question: str, text: str) -> float:
    question_terms = extract_keywords(question)
    if not question_terms:
        return 0.0

    text_terms = extract_keywords(text)
    if not text_terms:
        return 0.0

    return len(question_terms & text_terms) / len(question_terms)


def keyword_density(question: str, text: str) -> float:
    question_terms = extract_keywords(question)
    if not question_terms:
        return 0.0

    text_terms = [token for token in WORD_PATTERN.findall((text or "").lower()) if len(token) > 2]
    if not text_terms:
        return 0.0

    match_count = sum(1 for token in text_terms if token in question_terms)
    return min((match_count / max(len(text_terms), 1)) * 14, 1.0)


def is_heading_like(text: str) -> bool:
    normalized = normalize_text(text)
    words = normalized.split()
    if not words:
        return True

    if len(normalized) <= 60 and len(words) <= 8:
        return True

    if len(words) <= 10 and normalized.upper() == normalized:
        return True

    return not any(punctuation in normalized for punctuation in ".!?") and len(words) <= 6


def merge_short_paragraphs(paragraphs: list[str]) -> list[str]:
    merged: list[str] = []
    index = 0

    while index < len(paragraphs):
        current = normalize_text(paragraphs[index])
        if not current:
            index += 1
            continue

        while index + 1 < len(paragraphs):
            next_paragraph = normalize_text(paragraphs[index + 1])
            if not next_paragraph:
                index += 1
                continue

            combined = normalize_text(f"{current}\n{next_paragraph}")
            combined_words = len(combined.split())
            if is_heading_like(current):
                current = combined
                index += 1
                break

            if len(current.split()) < SHORT_PARAGRAPH_WORDS and combined_words <= TARGET_CHUNK_TOKENS:
                current = combined
                index += 1
                if len(current.split()) >= TARGET_CHUNK_TOKENS // 2:
                    break
                continue

            break

        merged.append(current)
        index += 1

    return merged


def coerce_optional_int(value: Any) -> Optional[int]:
    if value is None:
        return None

    if isinstance(value, int):
        return value

    text_value = str(value).strip()
    if not text_value:
        return None
    if text_value.isdigit():
        return int(text_value)

    match = re.search(r"(\d+)$", text_value)
    return int(match.group(1)) if match else None


def clamp_top_k(value: Any, default: int = DEFAULT_AGENT_TOP_K) -> int:
    if isinstance(value, bool):
        return default

    try:
        parsed = int(value)
    except (TypeError, ValueError):
        parsed = default

    return max(MIN_AGENT_TOP_K, min(MAX_AGENT_TOP_K, parsed))


def normalize_role(role: str) -> str:
    value = str(role or "").strip().lower()
    mapping = {
        ROLE_ADMIN.lower(): ROLE_ADMIN,
        ROLE_HR.lower(): ROLE_HR,
        ROLE_DEVELOPER.lower(): ROLE_DEVELOPER,
    }
    return mapping.get(value, str(role or "").strip())


def normalize_category_value(value: Any) -> str:
    return str(value or "").strip().upper()


def allowed_categories_for_role(role: str) -> list[str]:
    return ROLE_CATEGORY_ACCESS.get(normalize_role(role), [])


def allowed_visibility_scopes_for_role(role: str) -> list[str]:
    return ROLE_VISIBILITY_ACCESS.get(normalize_role(role), [])


def validate_category(category: str) -> str:
    normalized = normalize_category_value(category)
    if normalized not in VALID_CATEGORIES:
        raise ValueError(f"Invalid category '{category}'. Allowed values: {', '.join(VALID_CATEGORIES)}")
    return normalized


def normalize_visibility_scope(value: Any) -> str:
    normalized = str(value or "").strip().lower()
    return normalized if normalized in VALID_VISIBILITY_SCOPES else "private"


def validate_visibility_scope(scope: str) -> str:
    normalized = str(scope or "").strip().lower()
    if normalized not in VALID_VISIBILITY_SCOPES:
        raise ValueError(
            f"Invalid visibility scope '{scope}'. Allowed values: {', '.join(VALID_VISIBILITY_SCOPES)}"
        )
    return normalized


def category_matches_allowed(value: Any, allowed_categories: list[str]) -> bool:
    if not allowed_categories:
        return False
    return normalize_category_value(value) in {category.upper() for category in allowed_categories}


def visibility_matches_allowed(value: Any, allowed_scopes: list[str]) -> bool:
    if not allowed_scopes:
        return False
    return normalize_visibility_scope(value) in {scope.lower() for scope in allowed_scopes}


def build_category_where_filter(allowed_categories: list[str]) -> dict[str, Any]:
    category_values = sorted(
        {
            variant
            for category in allowed_categories
            for variant in CATEGORY_VARIANTS.get(category.upper(), {category.upper()})
        }
    )
    return {"category": {"$in": category_values}}


def build_access_where_filter(
    allowed_categories: list[str],
    allowed_visibility_scopes: list[str],
    document_id: str | None = None,
    document_ids: list[str] | None = None,
    page: int | None = None,
) -> dict[str, Any]:
    category_filter = build_category_where_filter(allowed_categories)
    visibility_filter = {"visibility_scope": {"$in": sorted({scope.lower() for scope in allowed_visibility_scopes})}}
    filters: list[dict[str, Any]] = [category_filter, visibility_filter]
    if document_ids:
        filters.append({"document_id": {"$in": sorted(set(document_ids))}})
    elif document_id:
        filters.append({"document_id": {"$eq": document_id}})
    if page is not None:
        filters.append({"page": {"$eq": page}})
    return {"$and": filters}


def format_confidence(score: float) -> str:
    percentage = max(0, min(99, round(score * 100)))
    return f"{percentage}%"


def build_metadata_payload(metadata: dict[str, Any]) -> dict[str, Any]:
    return {key: value for key, value in metadata.items() if value is not None}


def format_duration(seconds: float) -> str:
    if seconds < 60:
        return f"{seconds:.1f}s"

    minutes, remaining_seconds = divmod(seconds, 60)
    if minutes < 60:
        return f"{int(minutes)}m {remaining_seconds:04.1f}s"

    hours, remaining_minutes = divmod(int(minutes), 60)
    return f"{hours}h {remaining_minutes:02d}m {remaining_seconds:04.1f}s"


class EnterpriseRAGService:
    _chroma_client = None
    _embedder: Optional[SentenceTransformer] = None

    def __init__(self, collection_name: str = "enterprise_chunks", model_name: str = "llama-3.3-70b-versatile"):
        api_key = os.getenv("GROQ_API_KEY")
        if not api_key:
            raise ValueError("GROQ_API_KEY is missing. Add it to your .env file.")

        self.base_dir = Path(__file__).resolve().parent
        self.downloads_dir = self.base_dir / "downloads"
        self.downloads_dir.mkdir(parents=True, exist_ok=True)
        self.storage_dir = self.base_dir / "chroma_db_storage"
        self.collection_name = collection_name

        if EnterpriseRAGService._chroma_client is None:
            EnterpriseRAGService._chroma_client = chromadb.PersistentClient(path=str(self.storage_dir))

        self.client = EnterpriseRAGService._chroma_client
        self.collection = self.client.get_or_create_collection(self.collection_name)
        self.embedder = self._get_embedder()
        self.llm = ChatGroq(model=model_name, temperature=0, groq_api_key=api_key)
        self._document_chunk_cache: dict[str, list[dict[str, Any]]] = {}

    @classmethod
    def _get_embedder(cls) -> SentenceTransformer:
        if cls._embedder is None:
            cls._embedder = SentenceTransformer("sentence-transformers/all-mpnet-base-v2")
        return cls._embedder

    def _get_chroma_batch_size(self) -> int:
        try:
            batch_size = int(self.client.get_max_batch_size())
        except Exception:
            batch_size = INGEST_WRITE_BATCH_SIZE

        return max(batch_size, 1)

    def _iter_batches(self, items: list[Any], batch_size: int):
        safe_batch_size = max(batch_size, 1)
        for index in range(0, len(items), safe_batch_size):
            yield items[index : index + safe_batch_size]

    @staticmethod
    def _should_log_progress(current: int, total: int, interval: int) -> bool:
        if total <= 0:
            return True
        if current in {1, total}:
            return True
        return current % max(interval, 1) == 0

    @staticmethod
    def _build_progress_bar(current: int, total: int, width: int = 20) -> str:
        if total <= 0:
            return "[" + ("-" * width) + "]"

        ratio = max(0.0, min(current / total, 1.0))
        filled = min(width, max(0, round(ratio * width)))
        return "[" + ("#" * filled) + ("-" * (width - filled)) + "]"

    def _log_ingest(self, document_name: str, stage: str, message: str) -> None:
        timestamp = time.strftime("%H:%M:%S")
        print(f"[{timestamp}] [Ingest] [{document_name}] [{stage}] {message}", flush=True)

    def _log_verifier(self, query_text: str, stage: str, message: str) -> None:
        timestamp = time.strftime("%H:%M:%S")
        query_preview = normalize_text(query_text)[:96] or "empty-query"
        print(f"[{timestamp}] [Verifier] [{stage}] [{query_preview}] {message}", flush=True)

    @staticmethod
    def _coerce_bool(value: Any, default: bool = False) -> bool:
        if isinstance(value, bool):
            return value
        if isinstance(value, (int, float)):
            return bool(value)

        normalized = str(value or "").strip().lower()
        if normalized in {"true", "1", "yes"}:
            return True
        if normalized in {"false", "0", "no"}:
            return False
        return default

    def _log_ingest_progress(
        self,
        document_name: str,
        stage: str,
        current: int,
        total: int,
        unit_label: str,
        detail: str = "",
    ) -> None:
        progress_bar = self._build_progress_bar(current, total)
        percentage = 0 if total <= 0 else round((current / total) * 100)
        suffix = f" | {detail}" if detail else ""
        self._log_ingest(
            document_name,
            stage,
            f"{progress_bar} {percentage:>3}% ({current}/{total} {unit_label}){suffix}",
        )

    def save_upload(self, filename: str, file_bytes: bytes) -> Path:
        extension = Path(filename).suffix.lower()
        stored_path = self.downloads_dir / f"{uuid.uuid4().hex}{extension}"
        stored_path.write_bytes(file_bytes)
        return stored_path

    def _convert_to_pdf_if_needed(self, file_path: Path, document_name: str | None = None) -> Path:
        extension = file_path.suffix.lower()
        if extension == ".pdf":
            return file_path

        if extension not in {".ppt", ".pptx"}:
            raise ValueError(
                "Unsupported file type. Supported formats: PDF, PPT, PPTX, DOCX, TXT, MD, CSV, JSON, XLSX."
            )

        relative_base = Path("downloads") / file_path.stem
        if document_name:
            self._log_ingest(document_name, "CONVERT", "Converting presentation to PDF before ingestion.")
        converter = Ppt2Pdf(str(relative_base).replace("\\", "/"), extension[1:])
        converter.convert_ppt_to_pdf()
        pdf_path = self.base_dir / f"{relative_base}.pdf"

        if not pdf_path.exists():
            raise ValueError("PPT conversion failed. PDF output was not created.")

        return pdf_path

    def _extract_docx_pages(self, file_path: Path) -> list[dict[str, Any]]:
        document = DocxDocument(file_path)
        blocks: list[str] = []

        for paragraph in document.paragraphs:
            text = normalize_text(paragraph.text)
            if text:
                blocks.append(text)

        for table in document.tables:
            row_texts: list[str] = []
            for row in table.rows:
                cells = [normalize_text(cell.text) for cell in row.cells if normalize_text(cell.text)]
                if cells:
                    row_texts.append(" | ".join(cells))
            if row_texts:
                blocks.append("\n".join(row_texts))

        combined = "\n\n".join(blocks).strip()
        if not combined:
            return []

        return [{"page": 1, "text": normalize_structured_text(combined)}]

    def _extract_spreadsheet_pages(self, file_path: Path) -> list[dict[str, Any]]:
        workbook = load_workbook(filename=file_path, read_only=True, data_only=True)
        pages: list[dict[str, Any]] = []

        try:
            for sheet_index, sheet in enumerate(workbook.worksheets, start=1):
                rows: list[str] = []
                for row in sheet.iter_rows(values_only=True):
                    cells = [normalize_text(str(cell)) for cell in row if cell is not None and normalize_text(str(cell))]
                    if cells:
                        rows.append(" | ".join(cells))

                sheet_text = normalize_structured_text(f"Sheet: {sheet.title}\n\n" + "\n".join(rows))
                if sheet_text:
                    pages.append({"page": sheet_index, "text": sheet_text})
        finally:
            workbook.close()

        return pages

    def _extract_text_file_pages(self, file_path: Path) -> list[dict[str, Any]]:
        extension = file_path.suffix.lower()

        if extension == ".csv":
            with file_path.open("r", encoding="utf-8", errors="ignore", newline="") as handle:
                reader = csv.reader(handle)
                rows = [" | ".join(normalize_text(cell) for cell in row if normalize_text(cell)) for row in reader]
                text = "\n".join(row for row in rows if row.strip())
        elif extension == ".json":
            raw = file_path.read_text(encoding="utf-8", errors="ignore")
            try:
                parsed = json.loads(raw)
                text = json.dumps(parsed, indent=2, ensure_ascii=True)
            except json.JSONDecodeError:
                text = raw
        else:
            text = file_path.read_text(encoding="utf-8", errors="ignore")

        normalized = normalize_structured_text(text)
        if not normalized:
            return []

        return [{"page": 1, "text": normalized}]

    def _extract_pages_from_file(
        self,
        file_path: Path,
        document_name: str | None = None,
    ) -> list[dict[str, Any]]:
        extension = file_path.suffix.lower()

        if extension == ".pdf":
            return self._extract_pages(file_path, document_name=document_name)
        if extension in {".ppt", ".pptx"}:
            converted_pdf_path = self._convert_to_pdf_if_needed(file_path, document_name=document_name)
            return self._extract_pages(converted_pdf_path, document_name=document_name)
        if extension == ".docx":
            return self._extract_docx_pages(file_path)
        if extension == ".xlsx":
            return self._extract_spreadsheet_pages(file_path)
        if extension in {".txt", ".md", ".csv", ".json"}:
            return self._extract_text_file_pages(file_path)

        raise ValueError(
            "Unsupported file type. Supported formats: PDF, PPT, PPTX, DOCX, TXT, MD, CSV, JSON, XLSX."
        )

    def _extract_pages(self, pdf_path: Path, document_name: str | None = None) -> list[dict[str, Any]]:
        document = fitz.open(pdf_path)
        pages: list[dict[str, Any]] = []
        total_pages = len(document)

        if document_name:
            self._log_ingest(document_name, "EXTRACT", f"Reading {total_pages} PDF page(s).")

        try:
            for page_number, page in enumerate(document, start=1):
                blocks: list[str] = []
                for block in page.get_text("blocks", sort=True):
                    if len(block) < 5:
                        continue

                    block_text = normalize_structured_text(block[4])
                    if block_text:
                        blocks.append(block_text)

                page_text = "\n\n".join(blocks).strip()
                if page_text:
                    pages.append({"page": page_number, "text": page_text})

                if document_name and self._should_log_progress(
                    page_number,
                    total_pages,
                    INGEST_PROGRESS_PAGE_INTERVAL,
                ):
                    self._log_ingest_progress(
                        document_name,
                        "EXTRACT",
                        page_number,
                        total_pages,
                        "pages",
                        detail=f"{len(pages)} page(s) with extractable text",
                    )
        finally:
            document.close()

        return pages

    def _chunk_page(self, page: dict[str, Any]) -> list[dict[str, Any]]:
        chunks: list[dict[str, Any]] = []
        paragraphs = merge_short_paragraphs(split_paragraphs(page["text"]))
        for paragraph_index, paragraph in enumerate(paragraphs):
            tokens = normalize_text(" ".join(sentence_chunks(paragraph))).split()
            if not tokens:
                continue

            token_windows = [tokens] if len(tokens) <= TARGET_CHUNK_TOKENS else sliding_chunks(tokens)
            for window_index, token_window in enumerate(token_windows):
                chunk_text = normalize_text(" ".join(token_window))
                if len(chunk_text) >= 30:
                    chunks.append(
                        {
                            "page": page["page"],
                            "paragraph": paragraph_index,
                            "window": window_index,
                            "text": chunk_text,
                        }
                    )
        return chunks

    def _chunk_pages(
        self,
        pages: list[dict[str, Any]],
        document_name: str | None = None,
    ) -> list[dict[str, Any]]:
        chunks: list[dict[str, Any]] = []
        total_pages = len(pages)

        if document_name:
            self._log_ingest(document_name, "CHUNK", f"Building chunks from {total_pages} extracted page(s).")

        for page_index, page in enumerate(pages, start=1):
            page_chunks = self._chunk_page(page)
            chunks.extend(page_chunks)

            if document_name and self._should_log_progress(
                page_index,
                total_pages,
                INGEST_PROGRESS_PAGE_INTERVAL,
            ):
                self._log_ingest_progress(
                    document_name,
                    "CHUNK",
                    page_index,
                    total_pages,
                    "pages",
                    detail=f"{len(chunks)} chunk(s) prepared",
                )

        return chunks

    def _store_chunk_batch(
        self,
        *,
        document_id: str,
        document_name: str,
        category: str,
        sensitivity: str | None,
        visibility_scope: str,
        uploaded_by: str,
        chunk_batch: list[dict[str, Any]],
    ) -> int:
        texts = [chunk["text"] for chunk in chunk_batch]
        embeddings = self.embedder.encode(
            texts,
            batch_size=min(EMBEDDING_BATCH_SIZE, max(len(texts), 1)),
            show_progress_bar=False,
            normalize_embeddings=True,
        ).tolist()

        ids = [
            f"{document_id}:{chunk['page']}:{chunk['paragraph']}:{chunk['window']}"
            for chunk in chunk_batch
        ]
        metadatas = [
            build_metadata_payload(
                {
                    "chunk_id": f"{document_id}:{chunk['page']}:{chunk['paragraph']}:{chunk['window']}",
                    "document_id": document_id,
                    "document": document_name,
                    "category": category,
                    "sensitivity": sensitivity,
                    "visibility_scope": visibility_scope,
                    "page": chunk["page"],
                    "paragraph": chunk["paragraph"],
                    "window": chunk["window"],
                    "char_count": len(chunk["text"]),
                    "uploaded_by": uploaded_by,
                }
            )
            for chunk in chunk_batch
        ]

        self.collection.add(
            ids=ids,
            documents=texts,
            embeddings=embeddings,
            metadatas=metadatas,
        )
        return len(chunk_batch)

    def ingest_document(
        self,
        file_path: Path,
        document_name: str,
        category: str,
        sensitivity: str | None,
        visibility_scope: str,
        uploaded_by: str,
    ) -> IngestResult:
        validated_category = validate_category(category)
        validated_visibility_scope = validate_visibility_scope(visibility_scope)
        started_at = time.perf_counter()
        self._log_ingest(document_name, "START", f"Preparing ingestion for {file_path.name}.")
        pages = self._extract_pages_from_file(file_path, document_name=document_name)
        chunks = self._chunk_pages(pages, document_name=document_name)

        if not chunks:
            raise ValueError("No extractable text found in the uploaded document.")

        document_id = uuid.uuid4().hex
        total_chunks = len(chunks)
        total_batches = max(1, (total_chunks + INGEST_WRITE_BATCH_SIZE - 1) // INGEST_WRITE_BATCH_SIZE)
        indexed_chunks = 0
        self._log_ingest(
            document_name,
            "INDEX",
            f"Encoding and storing {total_chunks} chunk(s) in {total_batches} batch(es).",
        )

        try:
            for batch_index, offset in enumerate(range(0, total_chunks, INGEST_WRITE_BATCH_SIZE), start=1):
                batch_started_at = time.perf_counter()
                chunk_batch = chunks[offset : offset + INGEST_WRITE_BATCH_SIZE]
                indexed_chunks += self._store_chunk_batch(
                    document_id=document_id,
                    document_name=document_name,
                    category=validated_category,
                    sensitivity=sensitivity,
                    visibility_scope=validated_visibility_scope,
                    uploaded_by=uploaded_by,
                    chunk_batch=chunk_batch,
                )
                self._log_ingest_progress(
                    document_name,
                    "INDEX",
                    indexed_chunks,
                    total_chunks,
                    "chunks",
                    detail=(
                        f"batch {batch_index}/{total_batches} finished in "
                        f"{format_duration(time.perf_counter() - batch_started_at)}"
                    ),
                )
        except Exception:
            self._log_ingest(
                document_name,
                "ROLLBACK",
                "Indexing failed. Removing any partially stored chunks for this document.",
            )
            self.collection.delete(where={"document_id": {"$eq": document_id}})
            self._document_chunk_cache.pop(document_id, None)
            raise

        self._log_ingest(
            document_name,
            "DONE",
            (
                f"Finished ingestion: {len(pages)} page(s), {total_chunks} chunk(s), "
                f"{format_duration(time.perf_counter() - started_at)} total."
            ),
        )

        return IngestResult(
            document_id=document_id,
            document=document_name,
            category=validated_category,
            sensitivity=sensitivity,
            chunks_indexed=len(chunks),
        )

    @staticmethod
    def _select_top_chunks(
        chunks: list[RetrievedChunk],
        top_k: int,
        max_per_document: Optional[int] = None,
    ) -> list[RetrievedChunk]:
        selected: list[RetrievedChunk] = []
        seen: set[tuple[str, Optional[int], Optional[int], Optional[int], str]] = set()
        per_document_counts: dict[str, int] = {}

        for chunk in chunks:
            fingerprint = (
                chunk.document_id,
                chunk.page,
                chunk.paragraph,
                chunk.window,
                chunk.text[:180],
            )
            if fingerprint in seen:
                continue

            current_count = per_document_counts.get(chunk.document_id, 0)
            if max_per_document is not None and current_count >= max_per_document:
                continue

            selected.append(chunk)
            seen.add(fingerprint)
            per_document_counts[chunk.document_id] = current_count + 1
            if len(selected) >= top_k:
                break

        return selected

    def _rank_chunks(
        self,
        question: str,
        results: dict[str, Any],
        query_text: str,
    ) -> list[RetrievedChunk]:
        documents = (results.get("documents") or [[]])[0]
        metadatas = (results.get("metadatas") or [[]])[0]
        distances = (results.get("distances") or [[]])[0]

        ranked_chunks: list[RetrievedChunk] = []
        for index, (document_text, metadata, distance) in enumerate(zip(documents, metadatas, distances), start=1):
            if not document_text or not metadata:
                continue

            text = normalize_text(document_text)
            if not text:
                continue

            semantic_score = 1.0 / (1.0 + max(float(distance or 0.0), 0.0))
            lexical_score = keyword_overlap(question, text)
            density_score = keyword_density(question, text)
            combined_score = round((semantic_score * 0.64) + (lexical_score * 0.24) + (density_score * 0.12), 4)
            document_id = str(metadata.get("document_id") or metadata.get("doc_uuid") or "unknown")
            page = coerce_optional_int(metadata.get("page"))
            paragraph = coerce_optional_int(metadata.get("paragraph"))
            window = coerce_optional_int(metadata.get("window"))
            chunk_id = str(
                metadata.get("chunk_id")
                or f"{document_id}:{page or 'na'}:{paragraph or 0}:{window or 0}:{index}"
            )

            ranked_chunks.append(
                RetrievedChunk(
                    id=chunk_id,
                    document_id=document_id,
                    document=str(metadata.get("document") or "Unknown document"),
                    category=normalize_category_value(metadata.get("category") or "GENERAL"),
                    sensitivity=metadata.get("sensitivity"),
                    page=page,
                    paragraph=paragraph,
                    window=window,
                    text=text,
                    score=combined_score,
                    semantic_score=semantic_score,
                    lexical_score=lexical_score,
                    retrieval_query=query_text,
                )
            )

        ranked_chunks.sort(
            key=lambda item: (
                item.score,
                item.lexical_score,
                item.semantic_score,
            ),
            reverse=True,
        )
        return ranked_chunks

    def _load_document_chunks(self, document_id: str) -> list[dict[str, Any]]:
        if document_id in self._document_chunk_cache:
            return self._document_chunk_cache[document_id]

        results = self.collection.get(
            where={"document_id": {"$eq": document_id}},
            include=["documents", "metadatas"],
        )

        records: list[dict[str, Any]] = []
        for document_text, metadata in zip(results.get("documents") or [], results.get("metadatas") or []):
            if not document_text or not metadata:
                continue

            text = normalize_text(document_text)
            if not text:
                continue

            records.append(
                {
                    "text": text,
                    "page": coerce_optional_int(metadata.get("page")),
                    "paragraph": coerce_optional_int(metadata.get("paragraph")),
                    "window": coerce_optional_int(metadata.get("window")),
                }
            )

        records.sort(
            key=lambda item: (
                item["page"] or 0,
                item["paragraph"] or 0,
                item["window"] or 0,
            )
        )
        self._document_chunk_cache[document_id] = records
        return records

    @staticmethod
    def _extract_requested_page(query_text: str) -> int | None:
        match = PAGE_REFERENCE_PATTERN.search(query_text or "")
        if not match:
            return None

        try:
            page_number = int(match.group(1))
        except (TypeError, ValueError):
            return None

        return page_number if page_number > 0 else None

    @staticmethod
    def _combine_page_records(records: list[dict[str, Any]]) -> str:
        ordered_records = sorted(
            records,
            key=lambda item: (
                item.get("paragraph") or 0,
                item.get("window") or 0,
            ),
        )

        page_texts: list[str] = []
        for record in ordered_records:
            text = normalize_text(record.get("text") or "")
            if not text or text in page_texts:
                continue
            page_texts.append(text)

        return normalize_text(" ".join(page_texts))

    @staticmethod
    def _get_document_label(document_id: str) -> str:
        for record in list_document_records():
            if str(record.get("document_id") or "").strip() != document_id:
                continue
            document_name = str(record.get("document") or "").strip()
            if document_name:
                return document_name
        return document_id

    def _fetch_exact_page_chunks(
        self,
        query_text: str,
        role: str,
        document_id: str,
        page_number: int,
    ) -> list[RetrievedChunk]:
        results = self.collection.get(
            where=build_access_where_filter(
                allowed_categories_for_role(role),
                allowed_visibility_scopes_for_role(role),
                document_id=document_id,
                page=page_number,
            ),
            include=["documents", "metadatas"],
        )

        page_records: list[dict[str, Any]] = []
        document_name = self._get_document_label(document_id)
        category = "GENERAL"
        sensitivity = None

        for document_text, metadata in zip(results.get("documents") or [], results.get("metadatas") or []):
            if not document_text or not metadata:
                continue

            text = normalize_text(document_text)
            if not text:
                continue

            document_name = str(metadata.get("document") or document_name or document_id).strip() or document_id
            category = normalize_category_value(metadata.get("category") or category)
            sensitivity = metadata.get("sensitivity")
            page_records.append(
                {
                    "text": text,
                    "paragraph": coerce_optional_int(metadata.get("paragraph")),
                    "window": coerce_optional_int(metadata.get("window")),
                }
            )

        combined_text = self._combine_page_records(page_records)
        if not combined_text:
            return []

        overlap_score = keyword_overlap(query_text, combined_text)
        density_score = keyword_density(query_text, combined_text)
        exact_page_score = max(0.99, round((overlap_score * 0.5) + (density_score * 0.5), 4))

        return [
            RetrievedChunk(
                id=f"{document_id}:{page_number}:page",
                document_id=document_id,
                document=document_name,
                category=category,
                sensitivity=sensitivity,
                page=page_number,
                paragraph=0,
                window=0,
                text=combined_text,
                score=exact_page_score,
                semantic_score=exact_page_score,
                lexical_score=max(overlap_score, 0.5),
                retrieval_query=query_text,
            )
        ]

    def _expand_chunk_context(self, question: str, chunk: RetrievedChunk) -> RetrievedChunk:
        if not chunk.document_id or chunk.page is None:
            return chunk

        document_chunks = self._load_document_chunks(chunk.document_id)
        if not document_chunks:
            return chunk

        neighboring_texts: list[str] = []
        total_tokens = 0

        for record in document_chunks:
            if record["page"] != chunk.page:
                continue

            if chunk.paragraph is not None and record["paragraph"] is not None:
                if abs(record["paragraph"] - chunk.paragraph) > 1:
                    continue
                if (
                    record["paragraph"] == chunk.paragraph
                    and chunk.window is not None
                    and record["window"] is not None
                    and abs(record["window"] - chunk.window) > 1
                ):
                    continue
            elif neighboring_texts and total_tokens >= MAX_EXPANDED_CONTEXT_TOKENS:
                break

            record_text = record["text"]
            if not record_text or record_text in neighboring_texts:
                continue

            record_tokens = record_text.split()
            if neighboring_texts and total_tokens + len(record_tokens) > MAX_EXPANDED_CONTEXT_TOKENS:
                break

            neighboring_texts.append(record_text)
            total_tokens += len(record_tokens)

        expanded_text = normalize_text(" ".join(neighboring_texts))
        if not expanded_text:
            return chunk

        expanded_overlap = keyword_overlap(question, expanded_text)
        if len(expanded_text) <= len(chunk.text):
            return chunk

        if (
            expanded_overlap >= chunk.lexical_score
            or len(chunk.text) < 180
            or is_heading_like(chunk.text)
        ):
            density_score = keyword_density(question, expanded_text)
            updated_score = round(
                (chunk.semantic_score * 0.64) + (expanded_overlap * 0.24) + (density_score * 0.12),
                4,
            )
            return replace(
                chunk,
                text=expanded_text,
                lexical_score=expanded_overlap,
                score=updated_score,
            )

        return chunk

    @staticmethod
    def _extract_message_text(message: dict[str, Any]) -> str:
        content = message.get("content")
        if isinstance(content, str) and content.strip():
            return normalize_text(content)

        text = message.get("text")
        if isinstance(text, str) and text.strip():
            return normalize_text(text)

        return ""

    def _normalize_chat_history(self, chat_history: list[dict[str, Any]] | None) -> list[dict[str, str]]:
        normalized_history: list[dict[str, str]] = []
        for message in chat_history or []:
            role = str(message.get("role") or "").strip().lower()
            if role not in {"user", "assistant"}:
                continue

            text = self._extract_message_text(message)
            if not text:
                continue

            normalized_history.append({"role": role, "content": text})

        return normalized_history[-6:]

    @staticmethod
    def _history_to_text(chat_history: list[dict[str, str]]) -> str:
        if not chat_history:
            return "No prior conversation."

        lines = []
        for message in chat_history:
            speaker = "User" if message["role"] == "user" else "Assistant"
            lines.append(f"{speaker}: {message['content']}")
        return "\n".join(lines)

    @staticmethod
    def _document_record_sort_key(record: dict[str, Any]) -> tuple[str, str]:
        updated_at = str(record.get("updated_at") or "")
        created_at = str(record.get("created_at") or "")
        return updated_at, created_at

    def _get_active_document_records(self, role: str) -> list[dict[str, Any]]:
        allowed_categories = allowed_categories_for_role(role)
        allowed_visibility_scopes = allowed_visibility_scopes_for_role(role)
        records: list[dict[str, Any]] = []

        for record in list_document_records():
            document_id = str(record.get("document_id") or "").strip()
            if not document_id:
                continue
            if not category_matches_allowed(record.get("category"), allowed_categories):
                continue
            if not visibility_matches_allowed(record.get("visibility_scope"), allowed_visibility_scopes):
                continue
            records.append(record)

        return records

    def _resolve_fixed_document_id(
        self,
        query_text: str,
        role: str,
        requested_document_id: str | None = None,
    ) -> str | None:
        active_records = self._get_active_document_records(role)
        active_ids = {
            str(record.get("document_id") or "").strip()
            for record in active_records
            if str(record.get("document_id") or "").strip()
        }

        if requested_document_id:
            return requested_document_id if requested_document_id in active_ids else None

        normalized_query = normalize_text(query_text).lower()
        if not normalized_query:
            return None

        exact_matches: list[dict[str, Any]] = []
        stem_matches: list[dict[str, Any]] = []

        for record in active_records:
            document_name = str(record.get("document") or "").strip()
            if not document_name:
                continue

            full_name = document_name.lower()
            stem_name = Path(document_name).stem.lower()

            if full_name and full_name in normalized_query:
                exact_matches.append(record)
                continue

            if stem_name and len(stem_name) >= 4 and stem_name in normalized_query:
                stem_matches.append(record)

        matches = exact_matches or stem_matches
        if not matches:
            if not FILE_TYPE_HINT_PATTERN.search(normalized_query):
                return None

            query_terms = {
                token
                for token in extract_keywords(normalized_query)
                if len(token) >= 3 and not token.isdigit()
            }
            if not query_terms:
                return None

            scored_matches: list[tuple[tuple[int, int, str, str], dict[str, Any]]] = []
            for record in active_records:
                document_name = str(record.get("document") or "").strip()
                if not document_name:
                    continue

                document_terms = {
                    token
                    for token in WORD_PATTERN.findall(Path(document_name).stem.lower())
                    if len(token) >= 3 and token not in STOP_WORDS and not token.isdigit()
                }
                overlap = query_terms & document_terms
                if not overlap:
                    continue

                ranked_overlap = max(len(token) for token in overlap)
                scored_matches.append(
                    (
                        (
                            len(overlap),
                            ranked_overlap,
                            *self._document_record_sort_key(record),
                        ),
                        record,
                    )
                )

            if not scored_matches:
                return None

            scored_matches.sort(key=lambda item: item[0], reverse=True)
            best_score = scored_matches[0][0]
            tied_best = [record for score, record in scored_matches if score == best_score]
            if len(tied_best) != 1:
                return None
            matches = tied_best

        matches.sort(key=self._document_record_sort_key, reverse=True)
        resolved_id = str(matches[0].get("document_id") or "").strip()
        return resolved_id or None

    @staticmethod
    def _query_mentions_specific_document(query_text: str) -> bool:
        return bool(FILE_REFERENCE_PATTERN.search(query_text or ""))

    @staticmethod
    def _is_small_talk(query_text: str) -> bool:
        normalized = query_text.strip()
        return any(pattern.match(normalized) for pattern in SMALL_TALK_PATTERNS)

    @staticmethod
    def _looks_like_follow_up(query_text: str) -> bool:
        normalized = normalize_text(query_text).lower()
        if len(normalized.split()) <= 5:
            return True

        tokens = set(WORD_PATTERN.findall(normalized))
        return any(token in FOLLOW_UP_HINTS for token in tokens)

    def _rewrite_query_with_history(
        self,
        query_text: str,
        chat_history: list[dict[str, str]],
    ) -> str:
        if not chat_history:
            return query_text

        response = self.llm.invoke(
            [
                SystemMessage(
                    content=(
                        "Rewrite the user's latest question into a standalone enterprise search query. "
                        "Preserve intent, include missing references from the conversation, and return only the rewritten query. "
                        "If no rewrite is needed, return the original question."
                    )
                ),
                HumanMessage(
                    content=(
                        f"Conversation:\n{self._history_to_text(chat_history)}\n\n"
                        f"Latest question:\n{query_text}"
                    )
                ),
            ]
        )

        rewritten = normalize_text(self._read_llm_text(response.content))
        return rewritten or query_text

    def _generate_keyword_query(self, query_text: str) -> str:
        response = self.llm.invoke(
            [
                SystemMessage(
                    content=(
                        "Convert the user's request into a compact retrieval query for enterprise documents. "
                        "Keep only the core subject, constraints, and policy terms. Return one line only."
                    )
                ),
                HumanMessage(content=query_text),
            ]
        )
        keyword_query = normalize_text(self._read_llm_text(response.content))
        return keyword_query or query_text

    @staticmethod
    def _read_llm_text(content: Any) -> str:
        if isinstance(content, str):
            return content

        if isinstance(content, list):
            return "\n".join(
                item.get("text", "")
                for item in content
                if isinstance(item, dict) and item.get("type") == "text"
            )

        return ""

    def _plan_queries(
        self,
        query_text: str,
        chat_history: list[dict[str, str]],
    ) -> list[PlannedQuery]:
        planned_queries = [PlannedQuery(text=normalize_text(query_text), reason="original user question")]

        if chat_history and self._looks_like_follow_up(query_text):
            rewritten = self._rewrite_query_with_history(query_text, chat_history)
            if rewritten and rewritten.lower() != planned_queries[0].text.lower():
                planned_queries.append(PlannedQuery(text=rewritten, reason="rewritten follow-up with conversation context"))

        keyword_query = self._generate_keyword_query(planned_queries[0].text)
        if keyword_query and all(keyword_query.lower() != item.text.lower() for item in planned_queries):
            planned_queries.append(PlannedQuery(text=keyword_query, reason="compressed retrieval query"))

        return planned_queries

    @staticmethod
    def _heuristic_top_k(query_text: str, fallback: int = DEFAULT_AGENT_TOP_K) -> int:
        base_top_k = clamp_top_k(fallback, DEFAULT_AGENT_TOP_K)
        lowered = query_text.lower()

        if any(
            marker in lowered
            for marker in ("compare", "difference", "different", "list", "all", "summary", "summarize", "steps")
        ):
            return clamp_top_k(max(base_top_k, 6), base_top_k)

        if len(query_text.split()) <= 6:
            return clamp_top_k(max(MIN_AGENT_TOP_K, base_top_k - 1), base_top_k)

        return base_top_k

    @staticmethod
    def _extract_json_object(payload: str) -> dict[str, Any]:
        text = (payload or "").strip()
        fenced_match = re.search(r"```(?:json)?\s*(\{.*?\})\s*```", text, re.DOTALL)
        if fenced_match:
            text = fenced_match.group(1)

        start = text.find("{")
        end = text.rfind("}")
        if start == -1 or end == -1 or end < start:
            raise ValueError("No JSON object found in model output.")

        return json.loads(text[start : end + 1])

    def _parse_retrieval_action(
        self,
        payload: str,
        default_query: str,
        default_top_k: int,
        fixed_document_id: str | None = None,
    ) -> RetrievalAction:
        fallback = RetrievalAction(
            action="search",
            query=default_query,
            top_k=clamp_top_k(default_top_k, DEFAULT_AGENT_TOP_K),
            document_id=fixed_document_id,
            reason="fallback retrieval plan",
        )

        try:
            parsed = self._extract_json_object(payload)
        except (ValueError, json.JSONDecodeError, TypeError):
            return fallback

        action = str(parsed.get("action") or "search").strip().lower()
        if action not in {"search", "finish"}:
            action = "search"

        query = normalize_text(parsed.get("query") or default_query) or default_query
        document_id = fixed_document_id
        if not document_id:
            raw_document_id = str(parsed.get("document_id") or "").strip()
            document_id = raw_document_id or None
            if document_id and document_id.lower() in {"none", "null"}:
                document_id = None

        return RetrievalAction(
            action=action,
            query=query,
            top_k=clamp_top_k(parsed.get("top_k"), default_top_k),
            document_id=document_id,
            reason=normalize_text(parsed.get("reason") or "") or fallback.reason,
        )

    @staticmethod
    def _build_retrieval_observation(
        query_text: str,
        document_id: str | None,
        chunks: list[RetrievedChunk],
    ) -> str:
        if not chunks:
            scoped_document = document_id or "all accessible documents"
            return f'Search(query="{query_text}", document_id="{scoped_document}") returned no matches.'

        document_summaries: dict[str, dict[str, Any]] = {}
        for chunk in chunks:
            summary = document_summaries.setdefault(
                chunk.document_id,
                {"document": chunk.document, "count": 0, "best_score": 0.0},
            )
            summary["count"] += 1
            summary["best_score"] = max(summary["best_score"], chunk.score)

        lines = [
            f'Search(query="{query_text}", document_id="{document_id or "all"}") returned {len(chunks)} chunk(s).',
            "Document coverage: "
            + "; ".join(
                f'{value["document"]} [{doc_id}] x{value["count"]}, best={value["best_score"]:.3f}'
                for doc_id, value in document_summaries.items()
            ),
        ]

        for chunk in chunks[:6]:
            page_label = f"page {chunk.page}" if chunk.page else "page unavailable"
            lines.append(
                f'- {chunk.document} [{chunk.document_id}] | {page_label} | score {chunk.score:.3f} | {chunk.text[:160]}'
            )

        return "\n".join(lines)

    def _plan_retrieval_action(
        self,
        query_text: str,
        chat_history: list[dict[str, str]],
        observations: list[str],
        step: int,
        fixed_document_id: str | None,
        default_top_k: int,
    ) -> RetrievalAction:
        default_query = normalize_text(query_text) or query_text
        observation_block = "\n\n".join(observations[-2:]) if observations else "No retrieval observations yet."

        try:
            response = self.llm.invoke(
                [
                    SystemMessage(
                        content=(
                            "You control retrieval for an enterprise RAG system. "
                            "Pick one next action and return JSON only. "
                            "Allowed actions: "
                            "search = call the retrieval tool with a focused query and a top_k between 2 and 8; "
                            "finish = stop retrieving because the evidence is already sufficient. "
                            "Use lower top_k for narrow factual lookups, medium top_k for policy questions, and higher top_k for comparisons, lists, or ambiguous questions. "
                            "If one document clearly dominates, use its exact document_id to focus the next search. "
                            "Never invent a document_id. "
                            'Return JSON with keys: action, query, top_k, document_id, reason.'
                        )
                    ),
                    HumanMessage(
                        content=(
                            f"User question:\n{default_query}\n\n"
                            f"Recent conversation:\n{self._history_to_text(chat_history)}\n\n"
                            f"Current step: {step} of {MAX_AGENT_STEPS}\n"
                            f"Fixed document_id filter: {fixed_document_id or 'none'}\n\n"
                            f"Observations:\n{observation_block}"
                        )
                    ),
                ]
            )
            return self._parse_retrieval_action(
                payload=self._read_llm_text(response.content),
                default_query=default_query,
                default_top_k=default_top_k,
                fixed_document_id=fixed_document_id,
            )
        except Exception:
            return RetrievalAction(
                action="search",
                query=default_query,
                top_k=clamp_top_k(default_top_k, DEFAULT_AGENT_TOP_K),
                document_id=fixed_document_id,
                reason="fallback retrieval plan",
            )

    def _filter_results_by_access(
        self,
        results: dict[str, Any],
        allowed_categories: list[str],
        allowed_visibility_scopes: list[str],
        document_id: str | None = None,
        active_document_ids: set[str] | None = None,
    ) -> dict[str, list[list[Any]]]:
        documents = (results.get("documents") or [[]])[0]
        metadatas = (results.get("metadatas") or [[]])[0]
        distances = (results.get("distances") or [[]])[0]

        filtered_documents: list[Any] = []
        filtered_metadatas: list[Any] = []
        filtered_distances: list[Any] = []

        for document_text, metadata, distance in zip(documents, metadatas, distances):
            if not metadata:
                continue
            metadata_document_id = str(metadata.get("document_id") or "").strip()
            if document_id and metadata_document_id != document_id:
                continue
            if active_document_ids is not None and metadata_document_id not in active_document_ids:
                continue
            if not category_matches_allowed(metadata.get("category"), allowed_categories):
                continue
            if not visibility_matches_allowed(metadata.get("visibility_scope"), allowed_visibility_scopes):
                continue
            filtered_documents.append(document_text)
            filtered_metadatas.append(metadata)
            filtered_distances.append(distance)

        return {
            "documents": [filtered_documents],
            "metadatas": [filtered_metadatas],
            "distances": [filtered_distances],
        }

    def _query_enterprise_chunks(
        self,
        query_text: str,
        role: str,
        top_k: int,
        document_id: str | None = None,
        fetch_k: int | None = None,
        max_per_document: int | None = None,
    ) -> list[RetrievedChunk]:
        active_records = self._get_active_document_records(role)
        active_document_ids = {
            str(record.get("document_id") or "").strip()
            for record in active_records
            if str(record.get("document_id") or "").strip()
        }
        if not active_document_ids:
            return []

        if document_id and document_id not in active_document_ids:
            return []

        requested_top_k = clamp_top_k(top_k, DEFAULT_AGENT_TOP_K)
        query_embedding = self.embedder.encode(
            query_text,
            show_progress_bar=False,
            normalize_embeddings=True,
        ).tolist()

        effective_fetch_k = fetch_k or max(requested_top_k * 6, 24)
        filtered_results = self.collection.query(
            query_embeddings=[query_embedding],
            n_results=effective_fetch_k,
            where=build_access_where_filter(
                allowed_categories_for_role(role),
                allowed_visibility_scopes_for_role(role),
                document_id=document_id,
                document_ids=sorted(active_document_ids) if not document_id else None,
            ),
            include=["documents", "metadatas", "distances"],
        )
        ranked_chunks = self._rank_chunks(query_text, filtered_results, query_text=query_text)
        if ranked_chunks:
            return self._select_top_chunks(
                ranked_chunks,
                top_k=requested_top_k,
                max_per_document=max_per_document,
            )

        if self.collection.count() == 0:
            return []

        compatibility_results = self.collection.query(
            query_embeddings=[query_embedding],
            n_results=effective_fetch_k,
            include=["documents", "metadatas", "distances"],
        )
        locally_filtered = self._filter_results_by_access(
            compatibility_results,
            allowed_categories=allowed_categories_for_role(role),
            allowed_visibility_scopes=allowed_visibility_scopes_for_role(role),
            document_id=document_id,
            active_document_ids=active_document_ids,
        )
        ranked_chunks = self._rank_chunks(query_text, locally_filtered, query_text=query_text)
        return self._select_top_chunks(
            ranked_chunks,
            top_k=requested_top_k,
            max_per_document=max_per_document,
        )

    def update_document_visibility(self, document_id: str, visibility_scope: str) -> bool:
        validated_visibility_scope = validate_visibility_scope(visibility_scope)
        results = self.collection.get(
            where={"document_id": {"$eq": document_id}},
            include=["metadatas"],
        )

        ids = results.get("ids") or []
        metadatas = results.get("metadatas") or []
        if not ids:
            return False

        next_metadatas = [
            build_metadata_payload(
                {
                    **(metadata or {}),
                    "visibility_scope": validated_visibility_scope,
                }
            )
            for metadata in metadatas
        ]
        max_batch_size = self._get_chroma_batch_size()
        for id_batch, metadata_batch in zip(
            self._iter_batches(ids, max_batch_size),
            self._iter_batches(next_metadatas, max_batch_size),
        ):
            self.collection.update(ids=id_batch, metadatas=metadata_batch)
        return True

    def delete_document(self, document_id: str) -> bool:
        results = self.collection.get(
            where={"document_id": {"$eq": document_id}},
            include=[],
        )
        ids = results.get("ids") or []
        if ids:
            max_batch_size = self._get_chroma_batch_size()
            for id_batch in self._iter_batches(ids, max_batch_size):
                self.collection.delete(ids=id_batch)

        self._document_chunk_cache.pop(document_id, None)
        return bool(ids)

    def _query_legacy_chunks(
        self,
        query_text: str,
        chat_id: str | None,
        doc_uuid: str | None,
        top_k: int,
    ) -> list[RetrievedChunk]:
        if not chat_id:
            return []

        legacy_db = ChromaMultimodalDB(chat_id=chat_id, doc_uuid=doc_uuid)
        legacy_chunks = legacy_db.query_chunks(
            question=query_text,
            top_k=top_k,
            only_doc=doc_uuid,
            max_per_doc=2,
        )

        return [
            RetrievedChunk(
                id=chunk.source_id,
                document_id=chunk.doc_uuid,
                document=chunk.source_name,
                category="LEGACY",
                sensitivity=None,
                page=chunk.page,
                paragraph=chunk.paragraph,
                window=chunk.window,
                text=normalize_text(chunk.text),
                score=chunk.score,
                semantic_score=chunk.score,
                lexical_score=keyword_overlap(query_text, chunk.text),
                retrieval_query=query_text,
            )
            for chunk in legacy_chunks
            if normalize_text(chunk.text)
        ]

    @staticmethod
    def _merge_ranked_chunks(question: str, chunk_groups: list[list[RetrievedChunk]], top_k: int) -> list[RetrievedChunk]:
        scored: list[RetrievedChunk] = []
        seen: set[tuple[str, Optional[int], Optional[int], Optional[int], str]] = set()

        for chunk_group in chunk_groups:
            for chunk in chunk_group:
                fingerprint = (
                    chunk.document_id,
                    chunk.page,
                    chunk.paragraph,
                    chunk.window,
                    chunk.text[:180],
                )
                if fingerprint in seen:
                    continue
                seen.add(fingerprint)
                rescored = round((chunk.score * 0.82) + (keyword_overlap(question, chunk.text) * 0.18), 4)
                scored.append(replace(chunk, score=rescored))

        scored.sort(key=lambda item: item.score, reverse=True)
        return scored[:top_k]

    def _finalize_selected_chunks(
        self,
        question: str,
        chunks: list[RetrievedChunk],
        top_k: int,
    ) -> list[RetrievedChunk]:
        rescored: list[RetrievedChunk] = []

        for chunk in chunks:
            expanded_chunk = chunk if chunk.category == "LEGACY" else self._expand_chunk_context(question, chunk)
            overlap_score = keyword_overlap(question, expanded_chunk.text)
            density_score = keyword_density(question, expanded_chunk.text)
            updated_score = round(
                (expanded_chunk.semantic_score * 0.60) + (overlap_score * 0.26) + (density_score * 0.14),
                4,
            )
            rescored.append(
                replace(
                    expanded_chunk,
                    lexical_score=overlap_score,
                    score=updated_score,
                )
            )

        rescored.sort(
            key=lambda item: (
                item.score,
                item.lexical_score,
                item.semantic_score,
            ),
            reverse=True,
        )
        return self._select_top_chunks(rescored, top_k=top_k)

    def _run_agentic_enterprise_retrieval(
        self,
        query_text: str,
        role: str,
        chat_history: list[dict[str, str]],
        doc_uuid: str | None,
        top_k: int,
    ) -> list[RetrievedChunk]:
        desired_top_k = self._heuristic_top_k(query_text, top_k)
        observations: list[str] = []
        candidate_groups: list[list[RetrievedChunk]] = []

        for step in range(1, MAX_AGENT_STEPS + 1):
            action = self._plan_retrieval_action(
                query_text=query_text,
                chat_history=chat_history,
                observations=observations,
                step=step,
                fixed_document_id=doc_uuid,
                default_top_k=desired_top_k,
            )
            desired_top_k = clamp_top_k(action.top_k, desired_top_k)

            if action.action == "finish" and candidate_groups:
                break

            scoped_document_id = doc_uuid or action.document_id
            max_per_document = None if scoped_document_id else max(2, min(3, desired_top_k // 2 + 1))
            chunks = self._query_enterprise_chunks(
                query_text=action.query or query_text,
                role=role,
                top_k=desired_top_k,
                document_id=scoped_document_id,
                fetch_k=max(desired_top_k * 6, 24),
                max_per_document=max_per_document,
            )
            observations.append(
                self._build_retrieval_observation(
                    query_text=action.query or query_text,
                    document_id=scoped_document_id,
                    chunks=chunks,
                )
            )

            if not chunks:
                continue

            candidate_groups.append(chunks)
            merged_chunks = self._merge_ranked_chunks(
                question=query_text,
                chunk_groups=candidate_groups,
                top_k=desired_top_k,
            )
            if self._is_confident_enough(query_text, merged_chunks):
                return self._finalize_selected_chunks(query_text, merged_chunks, top_k=desired_top_k)

        if candidate_groups:
            merged_chunks = self._merge_ranked_chunks(
                question=query_text,
                chunk_groups=candidate_groups,
                top_k=desired_top_k,
            )
            return self._finalize_selected_chunks(query_text, merged_chunks, top_k=desired_top_k)

        planned_queries = self._plan_queries(query_text, chat_history)
        fallback_groups: list[list[RetrievedChunk]] = []
        for planned_query in planned_queries:
            chunks = self._query_enterprise_chunks(
                query_text=planned_query.text,
                role=role,
                top_k=desired_top_k,
                document_id=doc_uuid,
                fetch_k=max(desired_top_k * 6, 24),
                max_per_document=None if doc_uuid else max(2, min(3, desired_top_k // 2 + 1)),
            )
            if chunks:
                fallback_groups.append(chunks)

        if not fallback_groups:
            return []

        merged_chunks = self._merge_ranked_chunks(
            question=query_text,
            chunk_groups=fallback_groups,
            top_k=desired_top_k,
        )
        return self._finalize_selected_chunks(query_text, merged_chunks, top_k=desired_top_k)

    def _is_confident_enough(self, question: str, chunks: list[RetrievedChunk]) -> bool:
        if not chunks:
            return False

        top_chunk = chunks[0]
        top_overlap = keyword_overlap(question, top_chunk.text)

        if top_chunk.score >= MIN_HIGH_CONFIDENT_SCORE:
            return True

        if top_overlap >= MIN_LEXICAL_OVERLAP:
            return True

        if top_chunk.score >= MIN_CONFIDENT_SCORE and top_overlap >= MIN_LIGHT_OVERLAP:
            return True

        if len(chunks) >= 2:
            second_overlap = keyword_overlap(question, chunks[1].text)
            average_score = (top_chunk.score + chunks[1].score) / 2
            max_overlap = max(top_overlap, second_overlap)
            if average_score >= MIN_CONFIDENT_SCORE and max_overlap >= MIN_LIGHT_OVERLAP:
                return True

        return False

    def _build_context(self, chunks: list[RetrievedChunk]) -> str:
        lines: list[str] = []
        for index, chunk in enumerate(chunks, start=1):
            location = f"Page {chunk.page}" if chunk.page else "Page unavailable"
            lines.append(
                "\n".join(
                    [
                        f"[S{index}] Document: {chunk.document}",
                        f"Category: {chunk.category}",
                        f"Location: {location}",
                        f"Snippet: {chunk.text}",
                    ]
                )
            )
        return "\n\n".join(lines)

    def _build_answer_messages(self, query_text: str, chunks: list[RetrievedChunk]) -> list[Any]:
        context = self._build_context(chunks)
        return [
            SystemMessage(
                content=(
                    "You are an enterprise RAG assistant. "
                    "Answer strictly from the retrieved context. "
                    "If the context is insufficient, say so clearly. "
                    "Do not invent policy details."
                )
            ),
            HumanMessage(
                content=(
                    f"User question:\n{query_text}\n\n"
                    f"Authorized retrieved context:\n{context}\n\n"
                    "Write a concise answer grounded only in that context."
                )
            ),
        ]

    @staticmethod
    def _read_stream_text(content: Any) -> str:
        if isinstance(content, str):
            return _normalize_unicode(content)

        if isinstance(content, list):
            return "".join(
                item.get("text", "")
                for item in content
                if isinstance(item, dict) and item.get("type") == "text"
            )

        return ""

    def _generate_answer(self, query_text: str, chunks: list[RetrievedChunk]) -> str:
        response = self.llm.invoke(self._build_answer_messages(query_text, chunks))

        if isinstance(response.content, str):
            return normalize_text(response.content)

        if isinstance(response.content, list):
            text = "\n".join(
                item.get("text", "")
                for item in response.content
                if isinstance(item, dict) and item.get("type") == "text"
            )
            return normalize_text(text)

        return "No accessible data found for your role"

    def _stream_answer(self, query_text: str, chunks: list[RetrievedChunk]):
        for chunk in self.llm.stream(self._build_answer_messages(query_text, chunks)):
            delta = self._read_stream_text(getattr(chunk, "content", ""))
            if delta:
                yield delta

    def _answer_small_talk(self, query_text: str) -> QueryResult:
        response = self.llm.invoke(
            [
                SystemMessage(
                    content=(
                        "You are FindX, an enterprise document assistant. "
                        "Respond briefly to greetings or capability questions."
                    )
                ),
                HumanMessage(content=query_text),
            ]
        )
        answer = normalize_text(self._read_llm_text(response.content))
        return QueryResult(
            answer=answer or "Hello. I can help you search the indexed company documents you are allowed to access.",
            explanation="No document retrieval was needed because this request was conversational rather than document-specific.",
            sources=[],
        )

    def _build_explanation(self, chunks: list[RetrievedChunk]) -> str:
        if not chunks:
            return "No accessible data found for your role"

        unique_documents: list[str] = []
        for chunk in chunks:
            if chunk.document not in unique_documents:
                unique_documents.append(chunk.document)

        lead_chunk = chunks[0]
        document_summary = ", ".join(unique_documents[:3])
        if len(unique_documents) > 3:
            document_summary += ", and more"

        return (
            f"Grounded in {len(chunks)} retrieved passage(s) from {len(unique_documents)} document(s): "
            f"{document_summary}. Top supporting match scored about {format_confidence(lead_chunk.score)}."
        )

    def _build_sources(self, chunks: list[RetrievedChunk]) -> list[SourceItem]:
        return [
            SourceItem(
                id=chunk.id,
                document=chunk.document,
                doc_uuid=chunk.document_id,
                snippet=chunk.text[:480],
                page=chunk.page,
                confidence=format_confidence(chunk.score),
            )
            for chunk in chunks
        ]

    @staticmethod
    def _build_query_result_from_attempt(attempt: AnswerAttempt) -> QueryResult:
        return QueryResult(
            answer=attempt.answer,
            explanation=attempt.explanation,
            sources=attempt.sources,
        )

    @staticmethod
    def _stream_text_chunks(text: str, chunk_size: int = 32):
        normalized = text or ""
        for index in range(0, len(normalized), max(chunk_size, 1)):
            yield normalized[index : index + max(chunk_size, 1)]

    def _build_verifier_insufficiency_result(self, reason: str) -> QueryResult:
        explanation = normalize_text(reason) or (
            "Grounding verification could not confirm a safe answer from the accessible evidence."
        )
        return QueryResult(
            answer="I couldn't verify a grounded answer from the accessible evidence for this question.",
            explanation=explanation,
            sources=[],
        )

    def _build_answer_attempt(
        self,
        query_text: str,
        chunks: list[RetrievedChunk],
        refinement_count: int = 0,
    ) -> AnswerAttempt:
        return AnswerAttempt(
            answer=self._generate_answer(query_text, chunks),
            chunks=chunks,
            explanation=self._build_explanation(chunks),
            sources=self._build_sources(chunks),
            refinement_count=refinement_count,
        )

    def _build_verification_messages(
        self,
        context: ResolvedQueryContext,
        attempt: AnswerAttempt,
    ) -> list[Any]:
        scope_label = context.resolved_doc_uuid or "all accessible documents"
        exact_page_label = str(context.requested_page) if context.requested_page is not None else "none"
        authorized_context = self._build_context(attempt.chunks)

        return [
            SystemMessage(
                content=(
                    "You are a grounding verifier for an enterprise RAG system. "
                    "Check whether the draft answer is fully supported by the authorized context only. "
                    "Return JSON only with keys: grounded, needs_more_retrieval, gap_query, reason, keep_document_scope. "
                    "Set grounded=true only when all material claims in the draft answer are supported by the authorized context. "
                    "Set needs_more_retrieval=true only when one more focused retrieval pass is likely to fix the evidence gap. "
                    "Set gap_query to a compact retrieval query string or null. "
                    "Set keep_document_scope=true when any additional retrieval should stay inside the currently resolved document scope. "
                    "If the request targets an exact page and the current context from that page is insufficient, set grounded=false and needs_more_retrieval=false. "
                    "Do not answer in prose. Return JSON only."
                )
            ),
            HumanMessage(
                content=(
                    f"User question:\n{context.query_text}\n\n"
                    f"Resolved document scope:\n{scope_label}\n\n"
                    f"Exact page requested:\n{exact_page_label}\n\n"
                    f"Current refinement round:\n{attempt.refinement_count}\n\n"
                    f"Draft answer:\n{attempt.answer}\n\n"
                    f"Authorized context:\n{authorized_context}"
                )
            ),
        ]

    def _default_verification_result(
        self,
        query_text: str,
        chunks: list[RetrievedChunk],
        reason: str,
    ) -> VerificationResult:
        strong_enough = bool(chunks) and self._is_confident_enough(query_text, chunks)
        return VerificationResult(
            grounded=strong_enough,
            needs_more_retrieval=False,
            reason=normalize_text(reason) or "Verifier output could not be used safely.",
            gap_query=None,
        )

    def _parse_verification_result(
        self,
        payload: str,
        query_text: str,
        chunks: list[RetrievedChunk],
    ) -> VerificationResult:
        try:
            parsed = self._extract_json_object(payload)
        except (ValueError, json.JSONDecodeError, TypeError) as exc:
            return self._default_verification_result(
                query_text=query_text,
                chunks=chunks,
                reason=f"Verifier output could not be parsed safely: {exc}",
            )

        grounded = self._coerce_bool(parsed.get("grounded"), False)
        needs_more_retrieval = self._coerce_bool(parsed.get("needs_more_retrieval"), False)
        keep_document_scope = self._coerce_bool(parsed.get("keep_document_scope"), True)
        reason = normalize_text(parsed.get("reason") or "") or "Verifier did not provide a reason."
        gap_query_text = normalize_text(parsed.get("gap_query") or "")

        gap_query = None
        if not grounded and needs_more_retrieval and gap_query_text:
            gap_query = GapQuery(
                text=gap_query_text,
                reason=reason,
                keep_document_scope=keep_document_scope,
            )
        else:
            needs_more_retrieval = False

        if grounded:
            needs_more_retrieval = False
            gap_query = None

        return VerificationResult(
            grounded=grounded,
            needs_more_retrieval=needs_more_retrieval,
            reason=reason,
            gap_query=gap_query,
        )

    def _verify_answer_attempt(
        self,
        context: ResolvedQueryContext,
        attempt: AnswerAttempt,
    ) -> VerificationResult:
        self._log_verifier(
            context.query_text,
            "START",
            f"Verifying answer attempt {attempt.refinement_count + 1} with {len(attempt.chunks)} chunk(s).",
        )

        try:
            response = self.llm.invoke(self._build_verification_messages(context, attempt))
            payload = self._read_llm_text(response.content)
            verification = self._parse_verification_result(
                payload=payload,
                query_text=context.query_text,
                chunks=attempt.chunks,
            )
        except Exception as exc:
            verification = self._default_verification_result(
                query_text=context.query_text,
                chunks=attempt.chunks,
                reason=f"Verifier invocation failed: {exc}",
            )

        self._log_verifier(
            context.query_text,
            "RESULT",
            (
                f"grounded={verification.grounded} | "
                f"needs_more_retrieval={verification.needs_more_retrieval} | "
                f"reason={verification.reason}"
            ),
        )
        return verification

    def _prepare_query_context(
        self,
        query_text: str,
        role: str,
        chat_id: str | None,
        doc_uuid: str | None,
        chat_history: list[dict[str, Any]] | None,
        top_k: int,
    ) -> tuple[ResolvedQueryContext | None, QueryResult | None]:
        normalized_role = normalize_role(role)
        allowed_categories = allowed_categories_for_role(normalized_role)
        allowed_visibility_scopes = allowed_visibility_scopes_for_role(normalized_role)
        if not allowed_categories or not allowed_visibility_scopes:
            return None, QueryResult(
                answer="No accessible data found for your role",
                explanation="Your role does not have access to any searchable document categories.",
                sources=[],
            )

        if self._is_small_talk(query_text):
            return None, self._answer_small_talk(query_text)

        normalized_history = self._normalize_chat_history(chat_history)
        resolved_doc_uuid = self._resolve_fixed_document_id(
            query_text=query_text,
            role=normalized_role,
            requested_document_id=doc_uuid,
        )
        if (doc_uuid and not resolved_doc_uuid) or (
            not doc_uuid and self._query_mentions_specific_document(query_text) and not resolved_doc_uuid
        ):
            return None, QueryResult(
                answer="No accessible data found for your role",
                explanation="The request explicitly referenced a document that is deleted, unavailable, or outside your access scope.",
                sources=[],
            )

        return ResolvedQueryContext(
            query_text=query_text,
            role=normalized_role,
            chat_id=chat_id,
            normalized_history=normalized_history,
            resolved_doc_uuid=resolved_doc_uuid,
            requested_page=self._extract_requested_page(query_text),
            top_k=top_k,
        ), None

    def _retrieve_accessible_chunks(
        self,
        query_text: str,
        context: ResolvedQueryContext,
        document_scope: str | None,
    ) -> list[RetrievedChunk]:
        selected_chunks = self._run_agentic_enterprise_retrieval(
            query_text=query_text,
            role=context.role,
            chat_history=context.normalized_history,
            doc_uuid=document_scope,
            top_k=context.top_k,
        )

        if selected_chunks:
            return selected_chunks

        planned_queries = self._plan_queries(query_text, context.normalized_history)
        legacy_candidates: list[list[RetrievedChunk]] = []
        for planned_query in planned_queries:
            legacy_chunks = self._query_legacy_chunks(
                query_text=planned_query.text,
                chat_id=context.chat_id,
                doc_uuid=document_scope,
                top_k=context.top_k,
            )
            if legacy_chunks:
                legacy_candidates.append(legacy_chunks)

        if not legacy_candidates:
            return []

        return self._merge_ranked_chunks(
            question=query_text,
            chunk_groups=legacy_candidates,
            top_k=clamp_top_k(context.top_k, DEFAULT_AGENT_TOP_K),
        )

    def _resolve_query_chunks(
        self,
        context: ResolvedQueryContext,
    ) -> tuple[list[RetrievedChunk], QueryResult | None]:
        if context.requested_page is not None and context.resolved_doc_uuid:
            exact_page_chunks = self._fetch_exact_page_chunks(
                query_text=context.query_text,
                role=context.role,
                document_id=context.resolved_doc_uuid,
                page_number=context.requested_page,
            )
            if exact_page_chunks:
                return exact_page_chunks, None

            document_label = self._get_document_label(context.resolved_doc_uuid)
            return [], QueryResult(
                answer=(
                    f"No extractable content was found on page {context.requested_page} of {document_label}."
                ),
                explanation=(
                    "The request targeted an exact page in a specific accessible document, "
                    "but that page is not indexed or has no extractable text."
                ),
                sources=[],
            )

        selected_chunks = self._retrieve_accessible_chunks(
            query_text=context.query_text,
            context=context,
            document_scope=context.resolved_doc_uuid,
        )

        if not selected_chunks:
            return [], QueryResult(
                answer="No accessible data found for your role",
                explanation="The agent searched accessible documents, including rewritten retrieval attempts, but did not find a strong enough grounded match for this question.",
                sources=[],
            )

        return selected_chunks, None

    def _resolve_query(
        self,
        query_text: str,
        role: str,
        chat_id: str | None,
        doc_uuid: str | None,
        chat_history: list[dict[str, Any]] | None,
        top_k: int,
    ) -> tuple[list[RetrievedChunk], QueryResult | None]:
        context, early_result = self._prepare_query_context(
            query_text=query_text,
            role=role,
            chat_id=chat_id,
            doc_uuid=doc_uuid,
            chat_history=chat_history,
            top_k=top_k,
        )
        if early_result is not None or context is None:
            return [], early_result
        return self._resolve_query_chunks(context)

    def _run_verified_query(
        self,
        context: ResolvedQueryContext,
        selected_chunks: list[RetrievedChunk],
    ) -> QueryResult:
        attempt = self._build_answer_attempt(
            query_text=context.query_text,
            chunks=selected_chunks,
            refinement_count=0,
        )
        verification = self._verify_answer_attempt(context, attempt)
        attempt = replace(attempt, verification=verification)

        if verification.grounded and not verification.needs_more_retrieval:
            self._log_verifier(context.query_text, "FINAL", "Initial answer verified successfully.")
            return self._build_query_result_from_attempt(attempt)

        if (
            context.requested_page is None
            and verification.needs_more_retrieval
            and verification.gap_query is not None
            and attempt.refinement_count < MAX_VERIFICATION_REFINEMENTS
        ):
            refinement_scope = (
                context.resolved_doc_uuid
                if context.resolved_doc_uuid and verification.gap_query.keep_document_scope
                else None
            )
            scope_label = refinement_scope or "all accessible documents"
            self._log_verifier(
                context.query_text,
                "REFINE",
                f'Rerunning retrieval with "{verification.gap_query.text}" | scope={scope_label}',
            )
            refined_chunks = self._retrieve_accessible_chunks(
                query_text=verification.gap_query.text,
                context=context,
                document_scope=refinement_scope,
            )
            if refined_chunks:
                merged_chunks = self._merge_ranked_chunks(
                    question=context.query_text,
                    chunk_groups=[attempt.chunks, refined_chunks],
                    top_k=clamp_top_k(context.top_k, DEFAULT_AGENT_TOP_K),
                )
                merged_chunks = self._finalize_selected_chunks(
                    context.query_text,
                    merged_chunks,
                    top_k=clamp_top_k(context.top_k, DEFAULT_AGENT_TOP_K),
                )

                if merged_chunks:
                    refined_attempt = self._build_answer_attempt(
                        query_text=context.query_text,
                        chunks=merged_chunks,
                        refinement_count=attempt.refinement_count + 1,
                    )
                    refined_verification = self._verify_answer_attempt(context, refined_attempt)
                    refined_attempt = replace(refined_attempt, verification=refined_verification)

                    if refined_verification.grounded and not refined_verification.needs_more_retrieval:
                        self._log_verifier(
                            context.query_text,
                            "FINAL",
                            "Refined answer verified successfully.",
                        )
                        return self._build_query_result_from_attempt(refined_attempt)

                    self._log_verifier(
                        context.query_text,
                        "FINAL",
                        "Refined answer still failed grounding verification.",
                    )
                    return self._build_verifier_insufficiency_result(refined_verification.reason)

            self._log_verifier(
                context.query_text,
                "FINAL",
                "Verifier requested refinement but no stronger evidence was found.",
            )
            return self._build_verifier_insufficiency_result(verification.reason)

        self._log_verifier(
            context.query_text,
            "FINAL",
            "Grounding verification rejected the draft answer without a safe refinement path.",
        )
        return self._build_verifier_insufficiency_result(verification.reason)

    def query(
        self,
        query_text: str,
        role: str,
        chat_id: str | None = None,
        doc_uuid: str | None = None,
        chat_history: list[dict[str, Any]] | None = None,
        top_k: int = 4,
    ) -> QueryResult:
        context, early_result = self._prepare_query_context(
            query_text=query_text,
            role=role,
            chat_id=chat_id,
            doc_uuid=doc_uuid,
            chat_history=chat_history,
            top_k=top_k,
        )
        if early_result is not None:
            return early_result
        if context is None:
            return self._build_verifier_insufficiency_result(
                "The request could not be prepared for verified retrieval."
            )

        selected_chunks, chunk_error = self._resolve_query_chunks(context)
        if chunk_error is not None:
            return chunk_error

        return self._run_verified_query(context, selected_chunks)

    def stream_query(
        self,
        query_text: str,
        role: str,
        chat_id: str | None = None,
        doc_uuid: str | None = None,
        chat_history: list[dict[str, Any]] | None = None,
        top_k: int = 4,
    ):
        context, early_result = self._prepare_query_context(
            query_text=query_text,
            role=role,
            chat_id=chat_id,
            doc_uuid=doc_uuid,
            chat_history=chat_history,
            top_k=top_k,
        )

        if early_result is not None:
            yield {
                "type": "final",
                "answer": early_result.answer,
                "explanation": early_result.explanation,
                "sources": [source.model_dump() for source in early_result.sources],
            }
            return

        if context is None:
            fallback_result = self._build_verifier_insufficiency_result(
                "The request could not be prepared for verified retrieval."
            )
            yield {
                "type": "final",
                "answer": fallback_result.answer,
                "explanation": fallback_result.explanation,
                "sources": [source.model_dump() for source in fallback_result.sources],
            }
            return

        selected_chunks, chunk_error = self._resolve_query_chunks(context)
        if chunk_error is not None:
            yield {
                "type": "final",
                "answer": chunk_error.answer,
                "explanation": chunk_error.explanation,
                "sources": [source.model_dump() for source in chunk_error.sources],
            }
            return

        verified_result = self._run_verified_query(context, selected_chunks)
        for delta in self._stream_text_chunks(verified_result.answer):
            if delta:
                yield {"type": "token", "delta": delta}

        yield {
            "type": "final",
            "answer": verified_result.answer,
            "explanation": verified_result.explanation,
            "sources": [source.model_dump() for source in verified_result.sources],
        }
